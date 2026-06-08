"""
RAG 知识库问答系统 - LangChain 原生版本
支持 LangSmith 自动追踪
"""

import streamlit as st
import os
import json
from datetime import datetime
from dotenv import load_dotenv
import hashlib
import re
from typing import List, Tuple




# 加载环境变量
load_dotenv(override=True)

# ========== LangChain 导入 ==========
from langchain_deepseek import ChatDeepSeek
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough, RunnableLambda
from langchain_core.documents import Document
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings

# ========== 配置 ==========
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

# 对话保存目录
HISTORY_DIR = "chat_history"
if not os.path.exists(HISTORY_DIR):
    os.makedirs(HISTORY_DIR)


# ========== 初始化向量数据库（使用 LangChain 的 Chroma 封装） ==========
@st.cache_resource
def init_vector_store():
    """初始化 Chroma 向量数据库"""
    # 换用更轻量的模型（80MB vs 420MB，加载速度快 5-10 倍）
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",  # 改这里
        model_kwargs={'device': 'cpu'},
        encode_kwargs={'normalize_embeddings': False}
    )

    vectorstore = Chroma(
        collection_name="knowledge_base",
        embedding_function=embeddings,
        persist_directory="./chroma_db_langchain"
    )
    return vectorstore, embeddings


@st.cache_resource
def init_models():
    """初始化 LLM 模型"""
    model = ChatDeepSeek(
        model="deepseek-chat",
        temperature=0.3,
        max_tokens=1000
    )
    return model


# 初始化
vectorstore, embeddings = init_vector_store()
model = init_models()

# 创建检索器
retriever = vectorstore.as_retriever(search_kwargs={"k": 5})


# ========== 文档处理函数 ==========
def extract_text_from_pdf(file) -> str:
    import PyPDF2
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


def extract_text_from_txt(file) -> str:
    return file.read().decode("utf-8")


def extract_text_from_docx(file) -> str:
    import docx
    doc = docx.Document(file)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text


def split_text(text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> List[str]:
    """智能文本分块"""
    chunks = []
    start = 0
    text_len = len(text)
    while start < text_len:
        end = min(start + chunk_size, text_len)
        if end < text_len:
            for sep in ['。', '！', '？', '\n', '；']:
                pos = text.rfind(sep, start, end)
                if pos > start:
                    end = pos + 1
                    break
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - chunk_overlap if end < text_len else end
    return chunks


def add_document(filename: str, content: str) -> int:
    """添加文档到向量数据库"""
    chunks = split_text(content)

    # 创建 Document 对象
    documents = [
        Document(
            page_content=chunk,
            metadata={"source": filename, "chunk_index": i, "total_chunks": len(chunks)}
        )
        for i, chunk in enumerate(chunks)
    ]

    # 添加到向量存储
    vectorstore.add_documents(documents)

    return len(chunks)


def add_documents_batch(files) -> list:
    """批量添加文档"""
    results = []
    for file in files:
        try:
            if file.type == "application/pdf":
                text = extract_text_from_pdf(file)
            elif file.type == "text/plain":
                text = extract_text_from_txt(file)
            elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = extract_text_from_docx(file)
            else:
                results.append({"filename": file.name, "success": False, "error": "不支持的文件类型"})
                continue

            if text and text.strip():
                chunk_count = add_document(file.name, text)
                results.append({"filename": file.name, "success": True, "chunk_count": chunk_count, "size": len(text)})
            else:
                results.append({"filename": file.name, "success": False, "error": "文档内容为空"})
        except Exception as e:
            results.append({"filename": file.name, "success": False, "error": str(e)[:100]})
    return results


def get_document_list() -> list:
    """获取所有文档列表（从 metadata 中提取）"""
    try:
        # 获取所有文档的 metadata
        all_docs = vectorstore.get(include=["metadatas"])
        sources = set()
        for meta in all_docs.get('metadatas', []):
            if meta and 'source' in meta:
                sources.add(meta['source'])
        return sorted(list(sources))
    except Exception as e:
        return []


def delete_document(filename: str) -> bool:
    """删除指定文档"""
    try:
        # 获取要删除的文档 IDs
        all_docs = vectorstore.get(include=["metadatas"])
        ids_to_delete = []
        for i, meta in enumerate(all_docs.get('metadatas', [])):
            if meta and meta.get('source') == filename:
                ids_to_delete.append(all_docs['ids'][i])

        if ids_to_delete:
            vectorstore.delete(ids=ids_to_delete)
        return True
    except Exception as e:
        return False


def delete_all_documents() -> bool:
    """删除所有文档"""
    try:
        all_docs = vectorstore.get()
        if all_docs['ids']:
            vectorstore.delete(ids=all_docs['ids'])
        return True
    except Exception as e:
        return False


def get_document_count() -> int:
    """获取文档数量"""
    return len(get_document_list())


def get_chunk_count() -> int:
    """获取文本块数量"""
    try:
        all_docs = vectorstore.get()
        return len(all_docs.get('ids', []))
    except:
        return 0


# ========== 创建 RAG 链（LangSmith 可追踪） ==========
def format_docs(docs: List[Document]) -> str:
    """格式化检索到的文档"""
    if not docs:
        return "暂无相关文档"
    return "\n\n".join([doc.page_content for doc in docs])


def extract_sources(docs: List[Document]) -> set:
    """提取文档来源"""
    sources = set()
    for doc in docs:
        if doc.metadata and 'source' in doc.metadata:
            sources.add(doc.metadata['source'])
    return sources


# 创建提示词模板
prompt = ChatPromptTemplate.from_messages([
    ("system", """你是一个专业的文档问答助手。
你的任务是基于提供的文档内容回答用户问题。

要求：
1. 严格基于文档内容回答，不要编造信息
2. 如果文档中没有相关信息，直接说"文档中没有提到"
3. 回答要简洁、准确
4. 如果引用了文档，在最后注明来源"""),
    ("human", """文档内容：
{context}

用户问题：{question}""")
])


# 创建后处理函数
def post_process(response: str, docs: List[Document]) -> str:
    """后处理：添加来源信息"""
    sources = extract_sources(docs)
    if sources and "没有提到" not in response:
        response += f"\n\n---\n📖 参考文档：{', '.join(sources)}"
    return response


# 构建 RAG 链（这是 LangSmith 自动追踪的关键）
rag_chain = (
        {
            "context": retriever | RunnableLambda(format_docs),
            "question": RunnablePassthrough()
        }
        | prompt
        | model
        | StrOutputParser()
)


def rag_query(question: str, highlight_words: list = None) -> Tuple[str, List[dict], str]:
    """执行 RAG 查询（包装函数，保持与原接口兼容）"""
    # 先检索文档（用于返回片段信息）
    docs = retriever.invoke(question)

    # 执行 RAG 链
    answer = rag_chain.invoke(question)

    # 后处理添加来源
    sources = extract_sources(docs)
    if sources and "没有提到" not in answer:
        answer += f"\n\n---\n📖 参考文档：{', '.join(sources)}"

    # 构建片段信息
    chunks_info = []
    for i, doc in enumerate(docs[:3]):
        chunks_info.append({
            "content": doc.page_content,
            "source": doc.metadata.get('source', '未知'),
            "similarity": 1.0  # Chroma 不直接返回距离，可以后续优化
        })

    # 关键词高亮
    highlighted_answer = answer
    if highlight_words:
        for word in highlight_words:
            if word and len(word) > 1:
                pattern = re.compile(f'({re.escape(word)})', re.IGNORECASE)
                highlighted_answer = pattern.sub(r'<mark style="background-color: #ffff00;">\1</mark>',
                                                 highlighted_answer)

    return answer, chunks_info, highlighted_answer


# ========== 对话历史管理 ==========
def generate_conversation_signature(messages: list) -> str:
    if not messages:
        return ""
    key_messages = []
    for msg in messages:
        if msg["role"] in ["user", "assistant"]:
            content_hash = hashlib.md5(msg["content"][:100].encode()).hexdigest()[:12]
            key_messages.append(f"{msg['role']}_{content_hash}")
    return hashlib.md5("_".join(key_messages).encode()).hexdigest()


def get_saved_conversation_signatures() -> set:
    if not os.path.exists(HISTORY_DIR):
        return set()
    signatures = set()
    for f in os.listdir(HISTORY_DIR):
        if f.endswith('.json'):
            try:
                with open(os.path.join(HISTORY_DIR, f), 'r', encoding='utf-8') as file:
                    data = json.load(file)
                    if 'signature' in data:
                        signatures.add(data['signature'])
            except:
                pass
    return signatures


def is_conversation_saved(messages: list) -> bool:
    if not messages:
        return False
    signature = generate_conversation_signature(messages)
    if not signature:
        return False
    return signature in get_saved_conversation_signatures()


def save_chat_history(messages: list, title: str = None) -> tuple:
    if not messages:
        return None, "没有对话内容可保存"
    if is_conversation_saved(messages):
        return None, "当前对话已经保存过了，不会重复保存"

    if not title:
        title = f"chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

    filepath = os.path.join(HISTORY_DIR, f"{title}.json")
    counter = 1
    while os.path.exists(filepath):
        filepath = os.path.join(HISTORY_DIR, f"{title}_{counter}.json")
        counter += 1

    signature = generate_conversation_signature(messages)
    data = {
        "title": title,
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "signature": signature,
        "message_count": len(messages),
        "messages": messages
    }
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return os.path.basename(filepath), "保存成功"


def load_chat_history(filename: str) -> list:
    filepath = os.path.join(HISTORY_DIR, filename)
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return data.get("messages", [])
    except:
        return []


def get_chat_history_list() -> list:
    if not os.path.exists(HISTORY_DIR):
        return []
    files = []
    for f in os.listdir(HISTORY_DIR):
        if f.endswith('.json'):
            try:
                with open(os.path.join(HISTORY_DIR, f), 'r', encoding='utf-8') as file:
                    data = json.load(file)
                    files.append({
                        "filename": f,
                        "title": data.get("title", f),
                        "created_at": data.get("created_at", "未知"),
                        "message_count": data.get("message_count", len(data.get("messages", [])))
                    })
            except:
                files.append({"filename": f, "title": f, "created_at": "未知", "message_count": 0})
    return sorted(files, key=lambda x: x["created_at"], reverse=True)


def delete_chat_history(filename: str) -> bool:
    try:
        os.remove(os.path.join(HISTORY_DIR, filename))
        return True
    except:
        return False


# ========== 文档摘要和高亮 ==========
def get_document_summary(text: str, max_length: int = 200) -> str:
    if len(text) <= max_length:
        return text
    summary = text[:max_length]
    last_period = max(summary.rfind('。'), summary.rfind('. '), summary.rfind('\n'))
    if last_period > max_length * 0.5:
        summary = summary[:last_period + 1]
    else:
        summary += "..."
    return summary


# ========== Streamlit UI ==========
st.set_page_config(page_title="RAG 知识库问答", page_icon="📚", layout="wide")

st.markdown("""
<style>
    .stChatMessage { padding: 1rem; border-radius: 0.5rem; margin-bottom: 1rem; }
    mark { background-color: #ffff00; padding: 0 2px; border-radius: 3px; }
    .success-text { color: #00a65a; font-weight: bold; }
    .error-text { color: #dd4b39; font-weight: bold; }
</style>
""", unsafe_allow_html=True)

st.title("📚 RAG 知识库问答系统")
st.caption("上传文档（PDF / TXT / DOCX），然后提问 | 支持批量上传、对话历史、关键词高亮")

# 初始化 session state
if "messages" not in st.session_state:
    st.session_state.messages = []
if "highlight_keywords" not in st.session_state:
    st.session_state.highlight_keywords = []
if "batch_upload_results" not in st.session_state:
    st.session_state.batch_upload_results = []

# ========== 侧边栏 ==========
with st.sidebar:
    st.header("📄 文档管理")

    upload_mode = st.radio("上传方式", ["单个文件", "批量上传"], horizontal=True)

    if upload_mode == "单个文件":
        uploaded_file = st.file_uploader("选择文件", type=["pdf", "txt", "docx"])
        if uploaded_file:
            with st.spinner("正在处理文档..."):
                try:
                    if uploaded_file.type == "application/pdf":
                        text = extract_text_from_pdf(uploaded_file)
                    elif uploaded_file.type == "text/plain":
                        text = extract_text_from_txt(uploaded_file)
                    else:
                        text = extract_text_from_docx(uploaded_file)

                    if text and text.strip():
                        chunk_count = add_document(uploaded_file.name, text)
                        summary = get_document_summary(text, 200)
                        st.success(f"✅ 已添加: {uploaded_file.name}")
                        st.caption(f"📊 {chunk_count} 个文本块 | {len(text):,} 字符")
                        with st.expander("📖 文档预览"):
                            st.text(summary)
                    else:
                        st.error("文档内容为空")
                except Exception as e:
                    st.error(f"处理失败: {str(e)[:100]}")

    else:
        uploaded_files = st.file_uploader("选择多个文件", type=["pdf", "txt", "docx"], accept_multiple_files=True)
        if uploaded_files:
            col1, col2 = st.columns(2)
            with col1:
                if st.button("📤 开始批量上传", use_container_width=True):
                    with st.spinner(f"正在处理 {len(uploaded_files)} 个文件..."):
                        results = add_documents_batch(uploaded_files)
                        st.session_state.batch_upload_results = results
                        st.rerun()
            with col2:
                if st.button("🗑️ 清空结果", use_container_width=True):
                    st.session_state.batch_upload_results = []
                    st.rerun()

            if st.session_state.batch_upload_results:
                success_count = sum(1 for r in st.session_state.batch_upload_results if r["success"])
                st.markdown(f"<span class='success-text'>✅ 成功: {success_count}</span>", unsafe_allow_html=True)
                st.markdown(
                    f"<span class='error-text'>❌ 失败: {len(st.session_state.batch_upload_results) - success_count}</span>",
                    unsafe_allow_html=True)

    st.markdown("---")

    # 文档库管理
    st.subheader("📋 文档库")
    docs = get_document_list()

    if docs:
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🗑️ 删除全部", use_container_width=True):
                if delete_all_documents():
                    st.success("已删除所有文档")
                    st.rerun()
        with col2:
            st.caption(f"共 {len(docs)} 个文档")

        for doc in docs:
            col1, col2 = st.columns([4, 1])
            with col1:
                st.write(f"📄 {doc}")
            with col2:
                if st.button("🗑️", key=f"del_{doc}"):
                    if delete_document(doc):
                        st.success(f"已删除 {doc}")
                        st.rerun()
        st.caption(f"共 {get_chunk_count()} 个文本块")
    else:
        st.info("暂无文档")

    st.markdown("---")

    # 关键词高亮
    st.subheader("🔍 关键词高亮")
    highlight_input = st.text_input("输入关键词（用逗号分隔）", value=",".join(st.session_state.highlight_keywords))
    if highlight_input:
        st.session_state.highlight_keywords = [k.strip() for k in highlight_input.split(",") if k.strip()]

    st.markdown("---")

    # 对话历史
    st.subheader("💬 对话历史")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("保存对话", use_container_width=True):
            if st.session_state.messages:
                filename, msg = save_chat_history(st.session_state.messages)
                if filename:
                    st.success(msg)
                else:
                    st.warning(msg)
            else:
                st.warning("暂无对话内容")
    with col2:
        if st.button("开启新对话", use_container_width=True):
            st.session_state.messages = []
            st.rerun()

    chat_histories = get_chat_history_list()
    if chat_histories:
        with st.expander("📜 历史记录"):
            for hist in chat_histories[:5]:
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.caption(f"**{hist['title']}**")
                    st.caption(f"{hist['created_at']}")
                with col2:
                    if st.button("加载", key=f"load_{hist['filename']}"):
                        st.session_state.messages = load_chat_history(hist['filename'])
                        st.rerun()
                    if st.button("🗑️", key=f"del_hist_{hist['filename']}"):
                        delete_chat_history(hist['filename'])
                        st.rerun()

# ========== 主界面 ==========
col1, col2, col3 = st.columns(3)
with col1:
    st.metric("📚 文档数量", len(get_document_list()))
with col2:
    st.metric("📄 文本块数", get_chunk_count())
with col3:
    st.metric("💬 对话轮数", len(st.session_state.messages) // 2)

st.markdown("---")
st.header("💬 智能问答")

# 显示聊天历史
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        if msg["role"] == "assistant" and "chunks" in msg:
            st.markdown(msg.get("highlighted", msg["content"]), unsafe_allow_html=True)
            if msg.get("chunks"):
                with st.expander("🔗 查看检索到的文档片段"):
                    for i, chunk_info in enumerate(msg["chunks"][:3]):
                        st.caption(f"**片段{i + 1}** - 来源: {chunk_info['source']}")
                        st.text(chunk_info['content'][:300] + "...")
        else:
            st.markdown(msg["content"])

# 用户输入
prompt = st.chat_input("输入问题...")

if prompt:
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    with st.chat_message("assistant"):
        with st.spinner("🔍 检索文档中..."):
            answer, chunks, highlighted = rag_query(prompt, st.session_state.highlight_keywords)

        st.markdown(highlighted or answer, unsafe_allow_html=True)

        if chunks:
            with st.expander("🔗 查看检索到的文档片段"):
                for i, chunk_info in enumerate(chunks[:3]):
                    st.caption(f"**片段{i + 1}** - 来源: {chunk_info['source']}")
                    st.text(chunk_info['content'][:300] + "...")

    st.session_state.messages.append({
        "role": "assistant",
        "content": answer,
        "highlighted": highlighted,
        "chunks": chunks
    })
    st.rerun()

st.markdown("---")
st.caption("⚡ 支持 PDF、TXT、DOCX | 基于 LangChain + DeepSeek + ChromaDB | 对话历史保存 | 关键词高亮")